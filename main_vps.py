import os
import io
import zipfile
import shutil
import copy
from fastapi import FastAPI, UploadFile, File
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import lxml.etree as ET


import re
import time
import json
from fastapi import FastAPI, UploadFile, File, Form
from pdf2docx import Converter
from PIL import Image, ImageFilter, ImageEnhance
from htmldocx import HtmlToDocx

from typing import List
import uuid


# Register namespaces globally for ET

W   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
A   = "http://schemas.openxmlformats.org/drawingml/2006/main"
PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"
R   = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
WP  = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
VML = "urn:schemas-microsoft-com:vml"
WPS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
A14 = "http://schemas.microsoft.com/office/drawing/2010/main"
A14 = "http://schemas.microsoft.com/office/drawing/2010/main"
RELS_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CONTENT_TYPES_NS = "http://schemas.openxmlformats.org/package/2006/content-types"



app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/ping")
def ping():
    return {"status": "ok"}

@app.get("/")
def read_root():
    return {"status": "ok", "message": "Infidibulador Backend running"}

# Updated Accent Color
ACCENT_COLOR = "953735" 
FONT_NAME = "Cambria"

# Namespace constants are now global and registered above
# (Deleting local repeats to avoid confusion)




# ─── Structural Fixes ────────────────────────────────────────────────────────

def safe_rmtree(path, retries=5, delay=0.5):
    """Safely remove a directory with retries for Windows permission issues."""
    if not os.path.exists(path): return
    for i in range(retries):
        try:
            shutil.rmtree(path)
            return
        except Exception as e:
            if i == retries - 1: print(f"Failed to remove {path}: {e}")
            time.sleep(delay)

def remove_section_breaks(doc):
    """
    Remove all inner section breaks by deleting sectPr from paragraphs.
    Keep only the global document sectPr.
    """
    for para in doc.paragraphs:
        pPr = para._element.get_or_add_pPr()
        sectPr = pPr.find(qn("w:sectPr"))
        if sectPr is not None:
            pPr.remove(sectPr)

def remove_empty_lines_after_titles(paragraphs):
    """
    Remove empty paragraphs that immediately follow a title.
    """
    i = 0
    while i < len(paragraphs) - 1:
        para = paragraphs[i]
        text = para.text.strip()
        
        # Title detection logic (same as detect_and_format_titles)
        is_title = False
        if text and text[0].isdigit():
            fw = text.split(" ")[0]
            if all(c.isdigit() or c == "." for c in fw):
                is_title = True
        if any(s in para.style.name for s in ["Heading", "Title", "Caption"]):
            is_title = True
        if text.isupper() and len(text) > 5:
            is_title = True
        if para.runs and all(r.bold for r in para.runs):
            is_title = True
            
        if is_title:
            # Check for multiple empty paragraphs following the title
            j = i + 1
            while j < len(paragraphs):
                next_para = paragraphs[j]
                
                # FIX: Protect paragraphs that contain images even if they have no text
                # Aggressive Search: DrawingML, VML, Legacy Shapes, Objects, and Multimedia
                tags = ['drawing', 'pict', 'object', 'image', 'movie', 'audio', 'video']
                xpath_parts = [f'local-name()="{t}"' for t in tags]
                xpath_parts.append('contains(local-name(), "shape")') # shape, shapetype, group, rect, etc.
                
                xpath_query = f".//*[{' or '.join(xpath_parts)}]"
                has_image = bool(next_para._element.xpath(xpath_query))
                
                if not next_para.text.strip() and not has_image:
                    p = next_para._element
                    p.getparent().remove(p)
                    paragraphs.pop(j)
                else:
                    break
        i += 1

# ─── Header / Footer ─────────────────────────────────────────────────────────

# normalize_style_name removed

def inject_headers_footers(tpl_path: str, temp_dir: str):
    tpl_dir = "_tpl_hf_tmp"
    if os.path.exists(tpl_dir): shutil.rmtree(tpl_dir)
    with zipfile.ZipFile(tpl_path, "r") as z: z.extractall(tpl_dir)

    tpl_word = os.path.join(tpl_dir, "word")
    tgt_word = os.path.join(temp_dir, "word")
    tgt_rels_dir = os.path.join(tgt_word, "_rels")
    tpl_rels_dir = os.path.join(tpl_word, "_rels")
    os.makedirs(tgt_rels_dir, exist_ok=True)

    # 1. document.xml.rels
    tpl_rels_path = os.path.join(tpl_rels_dir, "document.xml.rels")
    tgt_rels_path = os.path.join(tgt_rels_dir, "document.xml.rels")
    tgt_rels_tree = ET.parse(tgt_rels_path)
    tpl_rels_tree = ET.parse(tpl_rels_path)
    tgt_rels_root = tgt_rels_tree.getroot()
    tpl_rels_root = tpl_rels_tree.getroot()

    for r in list(tgt_rels_root):
        t = r.get("Type", "").lower()
        if "header" in t or "footer" in t: tgt_rels_root.remove(r)

    existing_nums = []
    for rel in tgt_rels_root:
        rid = rel.get("Id", "rId0")
        num = "".join(c for c in rid if c.isdigit())
        if num: existing_nums.append(int(num))
    next_num = max(existing_nums, default=50) + 1

    old_to_new = {}
    hf_files_copied = []
    for rel in tpl_rels_root:
        rtype = rel.get("Type", "")
        if "header" not in rtype.lower() and "footer" not in rtype.lower(): continue
        old_rid = rel.get("Id")
        target_val = rel.get("Target", "")
        fname = os.path.basename(target_val)
        src = os.path.join(tpl_word, fname)
        dst = os.path.join(tgt_word, fname)
        if os.path.exists(src):
            shutil.copy2(src, dst)
            hf_files_copied.append(fname)
        hf_rels_src = os.path.join(tpl_rels_dir, fname + ".rels")
        hf_rels_dst = os.path.join(tgt_rels_dir, fname + ".rels")
        if os.path.exists(hf_rels_src): shutil.copy2(hf_rels_src, hf_rels_dst)

        new_rid = f"rId{next_num}"
        next_num += 1
        old_to_new[old_rid] = new_rid
        new_rel = ET.SubElement(tgt_rels_root, f"{{{RELS_NS}}}Relationship")
        new_rel.set("Id", new_rid)
        new_rel.set("Type", rtype)
        new_rel.set("Target", target_val)

    tgt_rels_tree.write(tgt_rels_path, xml_declaration=True, encoding="UTF-8", standalone=True)

    # 2. Media & Styles Injection (Simplest way)
    for folder in ["media", "styles.xml", "fontTable.xml", "settings.xml"]:
        src = os.path.join(tpl_word, folder)
        dst = os.path.join(tgt_word, folder)
        if os.path.exists(src):
            if os.path.isdir(src):
                os.makedirs(dst, exist_ok=True)
                for f in os.listdir(src): shutil.copy2(os.path.join(src, f), os.path.join(dst, f))
            else:
                shutil.copy2(src, dst)

    # 3. Content_Types overrides
    ct_path = os.path.join(temp_dir, "[Content_Types].xml")
    if os.path.exists(ct_path):
        ct_tree = ET.parse(ct_path)
        ct_root = ct_tree.getroot()
        existing_parts = {child.get("PartName", "").lower() for child in ct_root}
        h_ct = "application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml"
        f_ct = "application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml"
        for fname in hf_files_copied:
            part_name = f"/word/{fname}"
            if part_name.lower() not in existing_parts:
                ct_type = h_ct if "header" in fname.lower() else f_ct
                override = ET.SubElement(ct_root, f"{{{CONTENT_TYPES_NS}}}Override")
                override.set("PartName", part_name)
                override.set("ContentType", ct_type)
        ensure_content_types(temp_dir)

    # 4. Patch ALL sectPr in document.xml
    tpl_doc_path = os.path.join(tpl_word, "document.xml")
    tgt_doc_path = os.path.join(tgt_word, "document.xml")
    tpl_doc_tree = ET.parse(tpl_doc_path)
    tgt_doc_tree = ET.parse(tgt_doc_path)
    tpl_sectPr = tpl_doc_tree.getroot().find(f".//{{{W}}}sectPr")
    tgt_body = tgt_doc_tree.getroot().find(f"{{{W}}}body")

    if tpl_sectPr is not None and tgt_body is not None:
        tgt_sectPrs = list(tgt_body.findall(f".//{{{W}}}sectPr"))
        if not tgt_sectPrs: tgt_sectPrs = [ET.SubElement(tgt_body, f"{{{W}}}sectPr")]
        for sectPr in tgt_sectPrs:
            for tag in [f"{{{W}}}headerReference", f"{{{W}}}footerReference"]:
                for el in list(sectPr.findall(tag)): sectPr.remove(el)
            refs_to_insert = []
            for child in tpl_sectPr:
                if f"{{{W}}}headerReference" in child.tag or f"{{{W}}}footerReference" in child.tag:
                    new_el = copy.deepcopy(child)
                    if new_el.get(f"{{{R}}}id") in old_to_new:
                        new_el.set(f"{{{R}}}id", old_to_new[new_el.get(f"{{{R}}}id")])
                        refs_to_insert.append(new_el)
            for i, ref in enumerate(refs_to_insert): sectPr.insert(i, ref)
        tgt_doc_tree.write(tgt_doc_path, xml_declaration=True, encoding="UTF-8", standalone=True)
    
    shutil.rmtree(tpl_dir)

# ─── Header/Footer Tuning ────────────────────────────────────────────────────

def tune_hf_xml(temp_dir: str, theme=None):
    """
    Surgically colors shapes and text in Headers, Footers and Document.xml (Cover).
    """
    hf_color = (theme.get('headerFooter') if theme else ACCENT_COLOR).replace("#", "")
    title_color = (theme.get('title') if theme else ACCENT_COLOR).replace("#", "")
    # Use internalTitle (Bordas) for tables, fallback to hf_color
    table_color = (theme.get('internalTitle') or theme.get('headerFooter') or ACCENT_COLOR).replace("#", "")


    word_dir = os.path.join(temp_dir, "word")
    targets = [f for f in os.listdir(word_dir) if f.startswith("header") or f.startswith("footer")]
    targets.append("document.xml")


    for fname in targets:
        xml_path = os.path.join(word_dir, fname)
        if not os.path.exists(xml_path): continue
        
        is_footer = "footer" in fname
        is_header = "header" in fname
        tree = ET.parse(xml_path)
        root = tree.getroot()
        changed = False

        # ── 1. TABLES (Borders) ──
        # Header tables follow hf_color and only top/bottom borders
        # Others follow Bordas (internalTitle) logic
        current_tbl_color = hf_color if is_header else table_color
        
        for tc in root.findall(f".//{{{W}}}tc"):
            tcPr = tc.find(f"{{{W}}}tcPr")
            if tcPr is None: tcPr = ET.SubElement(tc, f"{{{W}}}tcPr")
            tcBorders = tcPr.find(f"{{{W}}}tcBorders")
            if tcBorders is None: tcBorders = ET.SubElement(tcPr, f"{{{W}}}tcBorders")
            
            # Sides to configure
            all_sides = ["top", "left", "bottom", "right", "insideH", "insideV"]
            for side in all_sides:
                b = tcBorders.find(f"{{{W}}}{side}")
                if b is None: b = ET.SubElement(tcBorders, f"{{{W}}}{side}")
                
                if is_header and side in ("left", "right", "insideV", "insideH"):
                    # Header tables: No side/internal borders
                    b.set(f"{{{W}}}val", "nil")
                else:
                    # Visible borders
                    b.set(f"{{{W}}}val", "single")
                    b.set(f"{{{W}}}sz", "6") 
                    b.set(f"{{{W}}}color", current_tbl_color)
            changed = True




        for para in root.findall(f".//{{{W}}}p"):
            pPr = para.find(f"{{{W}}}pPr")
            if pPr is None:
                pPr = ET.Element(f"{{{W}}}pPr")
                para.insert(0, pPr)
            
            # ── 2. PARAGRAPH BORDERS ──
            pBdr = pPr.find(f"{{{W}}}pBdr")
            if pBdr is not None:
                p_bdr_color = hf_color if (is_header or is_footer) else title_color
                for side_el in pBdr:
                    side_el.set(f"{{{W}}}color", p_bdr_color)
                    changed = True


            # ── 3. FONT (Header/Footer ONLY) ──
            if is_header or is_footer:
                for run in para.findall(f".//{{{W}}}r"):
                    rPr = run.find(f"{{{W}}}rPr")
                    if rPr is None: rPr = ET.SubElement(run, f"{{{W}}}rPr")
                    
                    if is_header:
                        # Force Cambria 9pt for Header
                        rFonts = rPr.find(f"{{{W}}}rFonts")
                        if rFonts is None: rFonts = ET.SubElement(rPr, f"{{{W}}}rFonts")
                        for attr in ["asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme"]:
                            if f"{{{W}}}{attr}" in rFonts.attrib: del rFonts.attrib[f"{{{W}}}{attr}"]
                        rFonts.set(f"{{{W}}}ascii", FONT_NAME)
                        rFonts.set(f"{{{W}}}hAnsi", FONT_NAME)
                        rFonts.set(f"{{{W}}}cs", FONT_NAME)
                        rFonts.set(f"{{{W}}}eastAsia", FONT_NAME)
                        
                        sz = rPr.find(f"{{{W}}}sz")
                        if sz is None: sz = ET.SubElement(rPr, f"{{{W}}}sz")
                        sz.set(f"{{{W}}}val", "18") # 9pt
                        szCs = rPr.find(f"{{{W}}}szCs")
                        if szCs is None: szCs = ET.SubElement(rPr, f"{{{W}}}szCs")
                        szCs.set(f"{{{W}}}val", "18") 
                        
                        target_color = "000000" # Header text is ALWAYS black
                    else:

                        # Footer text is WHITE and "sem cambria 9" (no override)
                        target_color = "FFFFFF"

                    # Special Case: Page Numbering (PAGE field) -> Always Cambria 9, White
                    is_page_num = False
                    # Check instrText in this paragraph
                    for it in para.findall(f".//{{{W}}}instrText"):
                        if it.text and "PAGE" in it.text.upper():
                            is_page_num = True
                            break
                    # Check fldSimple in this paragraph
                    if not is_page_num:
                        for fs in para.findall(f".//{{{W}}}fldSimple"):
                            instr = fs.get(f"{{{W}}}instr")
                            if instr and "PAGE" in instr.upper():
                                is_page_num = True
                                break
                    
                    if is_page_num:
                        # Force Cambria 9 for page numbers even in footer
                        rFonts = rPr.find(f"{{{W}}}rFonts")
                        if rFonts is None: rFonts = ET.SubElement(rPr, f"{{{W}}}rFonts")
                        for attr in ["asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme"]:
                            if f"{{{W}}}{attr}" in rFonts.attrib: del rFonts.attrib[f"{{{W}}}{attr}"]
                        rFonts.set(f"{{{W}}}ascii", FONT_NAME)
                        rFonts.set(f"{{{W}}}hAnsi", FONT_NAME)
                        rFonts.set(f"{{{W}}}cs", FONT_NAME)
                        rFonts.set(f"{{{W}}}eastAsia", FONT_NAME)
                        
                        sz = rPr.find(f"{{{W}}}sz")
                        if sz is None: sz = ET.SubElement(rPr, f"{{{W}}}sz")
                        sz.set(f"{{{W}}}val", "18")
                        szCs = rPr.find(f"{{{W}}}szCs")
                        if szCs is None: szCs = ET.SubElement(rPr, f"{{{W}}}szCs")
                        szCs.set(f"{{{W}}}val", "18")
                        target_color = "FFFFFF"


                    clr = rPr.find(f"{{{W}}}color")
                    if clr is None: clr = ET.SubElement(rPr, f"{{{W}}}color")
                    clr.set(f"{{{W}}}val", target_color)
                    changed = True




            # ── 3.5 Spacing and Empty Lines (Footer ONLY) ──
            if is_footer:
                # Remove any spacing
                for sp in pPr.findall(f"{{{W}}}spacing"): pPr.remove(sp)
                sp = ET.SubElement(pPr, f"{{{W}}}spacing")
                sp.set(f"{{{W}}}before", "0")
                sp.set(f"{{{W}}}after", "0")
                sp.set(f"{{{W}}}line", "240")
                sp.set(f"{{{W}}}lineRule", "auto")

                changed = True

            # ── 4. SHAPES (VML) ──
            current_accent = hf_color if (is_header or is_footer) else title_color
            for shape in para.xpath(".//v:rect | .//v:oval | .//v:shape", namespaces={"v": VML}):
                shape.set("fillcolor", f"#{current_accent}")
                shape.set("strokecolor", f"#{current_accent}")
                fill = shape.find(f"{{{VML}}}fill")
                if fill is None: fill = ET.SubElement(shape, f"{{{VML}}}fill")
                fill.set("on", "t")
                fill.set("color", f"#{current_accent}")
                fill.set("type", "solid")
                
                for txbx in shape.findall(f"{{{VML}}}textbox"):
                    txbx.set("inset", "0,0.1cm,0,0") # 0.1cm top padding
                changed = True


                # Text inside shape -> White
                for r in shape.xpath(".//w:r", namespaces={"w": W}):
                    rpr = r.find(f"{{{W}}}rPr")
                    if rpr is None: rpr = ET.SubElement(r, f"{{{W}}}rPr")
                    clr = rpr.find(f"{{{W}}}color")
                    if clr is None: clr = ET.SubElement(rpr, f"{{{W}}}color")
                    clr.set(f"{{{W}}}val", "FFFFFF")

            # ── 5. DRAWINGML (SDR/WPS) ──
            current_accent = hf_color if (is_header or is_footer) else title_color
            for wps_wsp in para.xpath(".//wps:wsp", namespaces={"wps": WPS}):
                spPr = wps_wsp.find(f"{{{WPS}}}spPr")
                if spPr is not None:
                    for old_fill in list(spPr.findall(f"{{{A}}}solidFill")): spPr.remove(old_fill)
                    solidFill = ET.SubElement(spPr, f"{{{A}}}solidFill")
                    ET.SubElement(solidFill, f"{{{A}}}srgbClr", val=current_accent)
                    
                    ln = spPr.find(f"{{{A}}}ln")
                    if ln is not None: spPr.remove(ln)
                    ln = ET.SubElement(spPr, f"{{{A}}}ln")
                    ln.set("w", "12700")
                    sfill = ET.SubElement(ln, f"{{{A}}}solidFill")
                    ET.SubElement(sfill, f"{{{A}}}srgbClr", val=current_accent)

                    
                    # Text inside DrawingML shape -> White
                    txbx = wps_wsp.find(f"{{{WPS}}}txbx")
                    if txbx is not None:
                        for r in txbx.xpath(".//w:r", namespaces={"w": W}):
                            rpr = r.find(f"{{{W}}}rPr")
                            if rpr is None: rpr = ET.SubElement(r, f"{{{W}}}rPr")
                            clr = rpr.find(f"{{{W}}}color")
                            if clr is None: clr = ET.SubElement(rpr, f"{{{W}}}color")
                            clr.set(f"{{{W}}}val", "FFFFFF")
                    changed = True
                    
            # Set body properties for DrawingML text box inset if it exists
            for bodyPr in para.xpath(".//wps:bodyPr", namespaces={"wps": WPS}):
                bodyPr.set("tIns", "36000") # 0.1cm (36000 EMUs)
                bodyPr.set("bIns", "0")
                bodyPr.set("lIns", "0")
                bodyPr.set("rIns", "0")
                changed = True

        if is_footer:
            # Recursively remove empty paragraphs from footers
            def remove_empty_p(node):
                to_remove = []
                for child in node:
                    if child.tag == f"{{{W}}}p":
                        texts = "".join(t.text for t in child.findall(f".//{{{W}}}t") if t.text)
                        has_visuals = child.findall(f".//v:shape", namespaces={"v": VML}) or child.findall(f".//v:oval", namespaces={"v": VML}) or child.findall(f".//v:rect", namespaces={"v": VML}) or child.findall(f".//wps:wsp", namespaces={"wps": WPS}) or child.findall(f".//w:drawing", namespaces={"w": W}) or child.findall(f".//w:pict", namespaces={"w": W})
                        if not texts.strip() and not has_visuals:
                            to_remove.append(child)
                    else:
                        remove_empty_p(child)
                for child in to_remove:
                    node.remove(child)
                    nonlocal changed
                    changed = True
            
            remove_empty_p(root)

        if changed:
            tree.write(xml_path, xml_declaration=True, encoding="UTF-8", standalone=True)

# ─── Image borders ───────────────────────────────────────────────────────────

def apply_image_borders_and_centering(temp_dir: str, theme=None):
    """
    Acts as a 'Document Analyst' using a surgical Regex Rule.
    - Limits processing to visual XMLs to avoid breaking system files (styles.xml).
    - Uses non-greedy patterns with negative lookahead to prevent corruption.
    - Validates XML structure BEFORE saving.
    - Specifically protects the cover image in document.xml.
    """
    word_dir = os.path.join(temp_dir, "word")
    # Use internalTitle for image borders as requested (Títulos Internos / Bordas)
    COR_HEX = (theme.get('internalTitle') if theme else ACCENT_COLOR).replace("#", "")
    NOVA_BORDA = f'<a:ln w="12700"><a:solidFill><a:srgbClr val="{COR_HEX}"/></a:solidFill></a:ln>'


    for root_path, _, files in os.walk(word_dir):
        for f in files:
            if not f.endswith(".xml"): continue
            if not any(k in f for k in ["document", "chart", "header", "footer"]): continue
            
            xml_path = os.path.join(root_path, f)
            try:
                with open(xml_path, 'r', encoding='utf-8') as file:
                    content = file.read()
                
                original_content = content
                placeholders = {}

                # 1. Protect Cover Drawing (id=9001)
                if "document.xml" in f:
                    def hide_cover(match):
                        p = f"__CVR_{len(placeholders)}__"
                        placeholders[p] = match.group(0)
                        return p
                    content = re.sub(r'<(w:drawing|v:shape|v:group|v:rect|v:oval)[^>]*>.*?(?:id="9001"|pic_9001|name="Cover").*?</\1>', hide_cover, content, flags=re.DOTALL)

                # 2. Surgical Replacement (Negative Lookahead: (?!</a:ln>). )
                # First, auto-closed: <a:ln .../>
                content = re.sub(r'<a:ln\b[^>]*/>', NOVA_BORDA, content)
                # Then, open/close: <a:ln ...>...</a:ln>
                content = re.sub(r'<a:ln\b[^>]*>(?:(?!</a:ln>).)*?</a:ln>', NOVA_BORDA, content, flags=re.DOTALL)

                # 3. Restore Cover
                for p, orig in placeholders.items():
                    content = content.replace(p, orig)

                # 4. VALIDATION: Ensure regex didn't break the XML
                if content != original_content:
                    try:
                        ET.fromstring(content.encode('utf-8')) # Check if parsable
                        with open(xml_path, 'w', encoding='utf-8') as file:
                            file.write(content)
                    except Exception as ve:
                        print(f"Validation failed for {f}, skipping regex. Error: {ve}")

            except Exception as e:
                print(f"File error in {f}: {e}")
                continue

            # 5. Centering & VML (Safe via lxml after regex)
            try:
                tree = ET.parse(xml_path)
                root = tree.getroot()
                changed = False
                
                # VML Strokes
                for shape in root.xpath(".//*[local-name()='shape' or local-name()='rect' or local-name()='oval' or local-name()='image' or local-name()='line']", namespaces={"v": VML}):
                    sid = shape.get("id") or ""
                    if sid == "9001" or "pic_9001" in sid: continue
                    if shape.tag.split("}")[1] in ("textbox", "fill", "shadow", "path", "textpath"): continue
                    shape.set("strokecolor", f"#{COR_HEX}")
                    shape.set("strokeweight", "1pt")
                    shape.set("stroked", "t")
                    changed = True

                # Centering Para
                if any(k in f for k in ["document", "header", "footer"]):
                    for para in root.findall(f".//{{{W}}}p"):
                        if para.xpath(".//w:drawing | .//w:pict", namespaces={"w": W, "v": VML}):
                            pPr = para.find(f"{{{W}}}pPr")
                            if pPr is None:
                                pPr = ET.Element(f"{{{W}}}pPr")
                                para.insert(0, pPr)
                            jc = pPr.find(f"{{{W}}}jc")
                            if jc is None: jc = ET.SubElement(pPr, f"{{{W}}}jc")
                            jc.set(f"{{{W}}}val", "center")
                            changed = True

                # 6. Sharpness (Word XML method - DISABLED per user request for pixel method)
                # (Removed as it produced imperceptible or failed results)
                pass

                if changed:

                    tree.write(xml_path, xml_declaration=True, encoding="UTF-8", standalone=True)
            except Exception as e:
                print(f"ET Parse skip for {f}: {e}")
                pass

# ─── Table borders ───────────────────────────────────────────────────────────

def set_table_borders(table, color="953735"):
    sides = ["top", "left", "bottom", "right", "insideH", "insideV"]
    color = color.replace("#", "")
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn("w:tcBorders"))
            if tcBorders is not None: tcPr.remove(tcBorders)
            tcBorders = OxmlElement("w:tcBorders")
            for side in sides:
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "single")
                b.set(qn("w:sz"), "6") # 0.75pt
                b.set(qn("w:color"), color)
                b.set(qn("w:space"), "0")
                tcBorders.append(b)
            tcPr.append(tcBorders)


# ─── High-level styling ──────────────────────────────────────────────────────

def apply_global_styles(doc, theme=None, indent_flag=True):

    # Set default document font to Cambria 12pt
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(12)
    # Set 6pt spacing between paragraphs
    style.paragraph_format.space_before = Pt(6)
    style.paragraph_format.space_after = Pt(6)
    
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.gutter = Cm(0)
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)


    # Force Cambria and spacing
    for para in doc.paragraphs:
        # Paragraph level overrides
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.0
        
        # Skip font override if already formatted by title logic
        if para.runs and para.runs[0].font.color and para.runs[0].font.color.rgb:
            continue
        for run in para.runs:
            run.font.name = FONT_NAME
            
        # 1.25cm Indent Logic
        # Apply Cm(1.25) if enabled, otherwise clear it
        if indent_flag:
            para.paragraph_format.first_line_indent = Cm(1.25)
        else:
            para.paragraph_format.first_line_indent = Cm(0)

            
    # Apply Header/Footer color if available
    if theme and 'headerFooter' in theme:
        hf_color = theme['headerFooter'].replace("#", "")
        for section in doc.sections:
            for hf in [section.header, section.footer]:
                for para in hf.paragraphs:
                    for run in para.runs:
                        try:
                            run.font.color.rgb = RGBColor.from_string(hf_color)
                        except: pass

    # FORCE TABLE FONT SIZE TO 10PT
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.line_spacing = 1.0
                    for run in para.runs:
                        run.font.name = FONT_NAME
                        run.font.size = Pt(10)

def insert_chapter_title(doc, filename, theme=None):
    """
    Inserts "Capítulo X" at the very beginning of the document.
    X is extracted from filename (e.g., "01_Artigo.docx" -> 1).
    Style: Book Antiqua, 45pt, Italic, Color from 'chapter'.
    """
    chapter_color = (theme.get('chapterX', theme.get('chapter', "#000000")) if theme else "#000000").replace("#", "")
    
    # Extract number from filename
    match = re.search(r'\d+', filename)
    chapter_num = match.group(0).lstrip('0') if match else "1"
    if not chapter_num: chapter_num = "0"
    
    # Insert at the top
    p = doc.paragraphs[0].insert_paragraph_before("")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Capítulo {chapter_num}")
    
    run.font.name = "Book Antiqua"
    run.font.size = Pt(45)
    run.font.italic = True
    # run.font.bold = False (Default is False, removing explicit True)



    try:
        run.font.color.rgb = RGBColor.from_string(chapter_color)
    except: pass
    
    # Force Book Antiqua in XML (sometimes necessary for certain fonts)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Book Antiqua")
    rFonts.set(qn("w:hAnsi"), "Book Antiqua")
    rPr.append(rFonts)

def detect_and_format_titles(paragraphs, theme=None):

    title_color = (theme['title'] if theme else ACCENT_COLOR).replace("#", "")
    chapter_color = (theme['chapter'] if theme else ACCENT_COLOR).replace("#", "")
    author_color = (theme['author'] if theme else ACCENT_COLOR).replace("#", "")
    internal_color = (theme['internalTitle'] if theme else title_color).replace("#", "")
    abstract_color = (theme['abstract'] if theme else "#FFFFFF").replace("#", "")

    for para in paragraphs:
        text = para.text.strip()
        if not text: continue
        
        is_chapter = False
        is_title = False
        is_author = False
        is_internal = False
        is_abstract = False
        
        # Chapter detection (1., 2., etc or specific styles)
        if text[0].isdigit():
            fw = text.split(" ")[0]
            if all(c.isdigit() or c == "." for c in fw): 
                is_chapter = True
        
        style_name = para.style.name.lower()
        if "heading 1" in style_name or "title" in style_name: 
            is_title = True
        
        if "heading 2" in style_name:
            is_internal = True
            
        if "heading 3" in style_name or "subtitle" in style_name:
            is_author = True
            
        # Detect bold as internal title if not already categorized
        if not (is_chapter or is_title or is_author or is_internal):
            if para.runs and all(r.bold for r in para.runs) and len(text) < 100:
                is_internal = True

        # Abstract detection
        if text.upper().startswith("RESUMO") or text.upper().startswith("ABSTRACT"):
            is_abstract = True

        current_color = title_color
        if is_chapter: current_color = chapter_color
        if is_author: current_color = author_color
        if is_internal: current_color = internal_color
        if is_abstract: current_color = abstract_color
        
        if is_chapter or is_title or is_author or is_internal or is_abstract:
            # Titles NEVER have indentation
            para.paragraph_format.first_line_indent = Cm(0)
            # Skip if it's the newly inserted "Capítulo X" (it's already formatted)
            if "Capítulo" in text and len(text) < 15 and para.runs and para.runs[0].font.size == Pt(45):
                continue

            for run in para.runs:
                # Other titles remain bold, but Article Title (is_title) is NO BOLD
                run.bold = False if is_title else True

                try:
                    run.font.color.rgb = RGBColor.from_string(current_color)
                except: pass
                
                if is_title:
                    run.font.name = "Cambria"
                    run.font.size = Pt(20)
                    run.font.italic = True
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    para.paragraph_format.space_before = Pt(20)
                else:
                    run.font.name = FONT_NAME




def format_tables(doc, theme=None):
    # Use internalTitle for tables (Títulos Internos / Bordas) fallback to accent
    table_color = (theme.get('internalTitle') or theme.get('headerFooter') or ACCENT_COLOR).replace("#", "")


    for table in doc.tables:
        set_table_borders(table, table_color)
        for i, row in enumerate(table.rows):
            is_header = (i == 0)
            
            # Shading logic: Header uses theme, Data rows use zebra striping (#F2F2F2 for even indices)
            fill_color = None
            if is_header:
                fill_color = table_color
            elif i > 0 and (i % 2 == 0):
                fill_color = "F2F2F2"

            for cell in row.cells:
                if fill_color:
                    tcPr = cell._tc.get_or_add_tcPr()
                    # Clean up existing shading to avoid duplicates
                    for old in list(tcPr.findall(qn("w:shd"))): tcPr.remove(old)
                    
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), fill_color)
                    tcPr.append(shd)

                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = FONT_NAME
                        run.font.size = Pt(10)
                        if is_header:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        else:
                            run.font.color.rgb = RGBColor(0, 0, 0)



def apply_pixel_sharpness(temp_dir: str):
    """
    Applies real pixel-level sharpening to all images in the document's media folder.
    Except "_cover" images to follow the 'no cover' rule.
    """
    media_dir = os.path.join(temp_dir, "word", "media")
    if not os.path.exists(media_dir): return
    
    images = [f for f in os.listdir(media_dir)
              if f.lower().endswith(('.png', '.jpeg', '.jpg', '.gif', '.bmp', '.tiff'))]

    # Skip files with "_cover" in name (our internal convention for theme covers)
    for img_name in images:
        if "_cover" in img_name: continue
        
        path = os.path.join(media_dir, img_name)
        try:
            img = Image.open(path)
            orig_mode = img.mode
            
            # Use RGB for filters
            img_rgb = img.convert("RGB")
            
            # Pixel manipulation per requested 'skill': 
            # Step A: UnsharpMask (radius=2, percent=125, threshold=3)
            sharpened = img_rgb.filter(ImageFilter.UnsharpMask(radius=2, percent=125, threshold=3))
            
            # Step B: Additional 25% boost via ImageEnhance
            sharpened = ImageEnhance.Sharpness(sharpened).enhance(1.25)
            
            # Restore mode (alpha transparency)
            if orig_mode == "RGBA":
                alpha = img.split()[-1]
                sharpened = sharpened.convert("RGBA")
                sharpened.putalpha(alpha)
            elif orig_mode == "P":
                sharpened = sharpened.quantize(colors=256)
            elif orig_mode == "L":
                sharpened = sharpened.convert("L")
                
            fmt = img.format if img.format else ("PNG" if img_name.lower().endswith('.png') else "JPEG")
            sharpened.save(path, format=fmt, quality=95 if fmt == "JPEG" else None)

        except Exception as e:
            print(f"Error processing {img_name} for sharpness: {e}")

def extract_first_page_image(template_path: str):

    temp = "_tpl_img_tmp"
    if os.path.exists(temp): shutil.rmtree(temp)
    with zipfile.ZipFile(template_path, "r") as z: z.extractall(temp)
    doc_xml = os.path.join(temp, "word", "document.xml")
    rels_xml = os.path.join(temp, "word", "_rels", "document.xml.rels")
    if not os.path.exists(doc_xml): return None
    rId_map = {}
    if os.path.exists(rels_xml):
        tree = ET.parse(rels_xml)
        for rel in tree.getroot(): 
            rId_map[rel.get("Id")] = rel.get("Target", "")


    blips = ET.parse(doc_xml).getroot().findall(f".//{{{A}}}blip")
    if not blips: return None
    first_rid = blips[0].get(f"{{{R}}}embed")
    if not first_rid or first_rid not in rId_map: return None
    img_rel = rId_map[first_rid]
    img_full = os.path.join(temp, "word", img_rel)
    if not os.path.exists(img_full): return None
    dest = f"_cover_image{os.path.splitext(img_full)[1]}"
    shutil.copy2(img_full, dest)
    shutil.rmtree(temp)
    return dest

def ensure_content_types(temp_dir: str):
    ct_path = os.path.join(temp_dir, "[Content_Types].xml")
    media_dir = os.path.join(temp_dir, "word", "media")
    if not os.path.exists(ct_path): return
    
    try:
        tree = ET.parse(ct_path)
        root = tree.getroot()
        
        # Get all extensions currently in media/
        current_exts = set()
        if os.path.exists(media_dir):
            for f in os.listdir(media_dir):
                ext = os.path.splitext(f)[1].lower().lstrip('.')
                if ext: current_exts.add(ext)
        
        # Map of ext -> ContentType
        type_map = {
            "png": "image/png",
            "jpeg": "image/jpeg",
            "jpg": "image/jpeg",
            "gif": "image/gif",
            "bmp": "image/bmp",
            "tiff": "image/tiff",
            "tif": "image/tiff",
            "svg": "image/svg+xml"
        }
        
        # Find existing Default entries
        existing_defaults = {d.get("Extension").lower(): d for d in root.findall(f"{{{CONTENT_TYPES_NS}}}Default")}
        
        changed = False
        for ext in current_exts:
            if ext not in existing_defaults and ext in type_map:
                elem = ET.SubElement(root, f"{{{CONTENT_TYPES_NS}}}Default")
                elem.set("Extension", ext)
                elem.set("ContentType", type_map[ext])
                changed = True
        
        if changed:
            tree.write(ct_path, xml_declaration=True, encoding="UTF-8", standalone=True)
    except Exception as e:
        print(f"Error updating [Content_Types].xml: {e}")


def insert_cover_image(temp_dir: str, cover_image_path: str):
    print(f"DEBUG: [insert_cover_image] Start. Path: {cover_image_path}")
    if not os.path.exists(cover_image_path):
        print(f"DEBUG: [insert_cover_image] ERROR: File not found at {cover_image_path}")
        return

    media_dir = os.path.join(temp_dir, "word", "media")
    os.makedirs(media_dir, exist_ok=True)
    ext = os.path.splitext(cover_image_path)[1].lower().lstrip('.')
    dest_name = f"_cover.{ext}"
    shutil.copy2(cover_image_path, os.path.join(media_dir, dest_name))
    print(f"DEBUG: [insert_cover_image] Copied to {dest_name}")

    # Ensure content types are registered
    ensure_content_types(temp_dir)

    rels_path = os.path.join(temp_dir, "word", "_rels", "document.xml.rels")
    tree = ET.parse(rels_path)
    root = tree.getroot()
    
    # Get max rId
    nums = [int(re.search(r'\d+', r.get("Id")).group()) for r in root if re.search(r'\d+', r.get("Id") or "")]
    new_rid = f"rId{max(nums, default=0) + 1}"
    print(f"DEBUG: [insert_cover_image] New rId: {new_rid}")

    rel = ET.SubElement(root, f"{{{RELS_NS}}}Relationship")
    rel.set("Id", new_rid)
    rel.set("Type", f"{R}/image")
    rel.set("Target", f"media/{dest_name}")
    tree.write(rels_path, xml_declaration=True, encoding="UTF-8", standalone=True)

    doc_xml = os.path.join(temp_dir, "word", "document.xml")
    tree = ET.parse(doc_xml)
    root_el = tree.getroot()
    
    for bg in root_el.findall(f".//{{{W}}}background"): root_el.remove(bg)
    body = root_el.find(f"{{{W}}}body")
    if body is None:
        print("DEBUG: [insert_cover_image] ERROR: Body not found")
        return

    cx, cy = 7639200, 10749600
    off_x, off_y = -943200, -43200

    nsmap = {
        'w': W, 'wp': WP, 'a': A, 'pic': PIC, 'r': R, 'a14': A14
    }
    
    p = ET.Element(f"{{{W}}}p", nsmap=nsmap)
    pPr = ET.SubElement(p, f"{{{W}}}pPr")
    ET.SubElement(pPr, f"{{{W}}}jc").set(f"{{{W}}}val", "center")
    ET.SubElement(pPr, f"{{{W}}}spacing").attrib.update({f"{{{W}}}before": "0", f"{{{W}}}after": "0"})
    
    r_el = ET.SubElement(p, f"{{{W}}}r")
    drawing = ET.SubElement(r_el, f"{{{W}}}drawing")
    anchor = ET.SubElement(drawing, f"{{{WP}}}anchor")
    anchor.attrib.update({
        "distT": "0", "distB": "0", "distL": "114300", "distR": "114300",
        "simplePos": "0", "relativeHeight": "251658240", "behindDoc": "1",
        "locked": "0", "layoutInCell": "1", "allowOverlap": "1"
    })
    ET.SubElement(anchor, f"{{{WP}}}simplePos").attrib.update({"x": "0", "y": "0"})
    
    posH = ET.SubElement(anchor, f"{{{WP}}}positionH")
    posH.set("relativeFrom", "page")
    ET.SubElement(posH, f"{{{WP}}}posOffset").text = "0"
    
    posV = ET.SubElement(anchor, f"{{{WP}}}positionV")
    posV.set("relativeFrom", "page")
    ET.SubElement(posV, f"{{{WP}}}posOffset").text = "0"
    
    ET.SubElement(anchor, f"{{{WP}}}extent").attrib.update({"cx": str(cx), "cy": str(cy)})
    ET.SubElement(anchor, f"{{{WP}}}effectExtent").attrib.update({"l": "19050", "t": "0", "r": "0", "b": "0"})
    ET.SubElement(anchor, f"{{{WP}}}wrapNone")
    ET.SubElement(anchor, f"{{{WP}}}docPr").attrib.update({"id": "9001", "name": "Cover"})
    
    cNv = ET.SubElement(anchor, f"{{{WP}}}cNvGraphicFramePr")
    ET.SubElement(cNv, f"{{{A}}}graphicFrameLocks").set("noChangeAspect", "1")
    
    graphic = ET.SubElement(anchor, f"{{{A}}}graphic")
    data = ET.SubElement(graphic, f"{{{A}}}graphicData")
    data.set("uri", PIC)
    
    pic_el = ET.SubElement(data, f"{{{PIC}}}pic")
    nv = ET.SubElement(pic_el, f"{{{PIC}}}nvPicPr")
    ET.SubElement(nv, f"{{{PIC}}}cNvPr").attrib.update({"id": "9001", "name": "Cover"})
    ET.SubElement(ET.SubElement(nv, f"{{{PIC}}}cNvPicPr"), f"{{{A}}}picLocks").set("noChangeAspect", "1")
    
    blipFill = ET.SubElement(pic_el, f"{{{PIC}}}blipFill")
    blip = ET.SubElement(blipFill, f"{{{A}}}blip")
    blip.set(f"{{{R}}}embed", new_rid)
    
    stretch = ET.SubElement(blipFill, f"{{{A}}}stretch")
    ET.SubElement(stretch, f"{{{A}}}fillRect")
    
    spPr = ET.SubElement(pic_el, f"{{{PIC}}}spPr")
    xfrm = ET.SubElement(spPr, f"{{{A}}}xfrm")
    ET.SubElement(xfrm, f"{{{A}}}off").attrib.update({"x": "0", "y": "0"})
    ET.SubElement(xfrm, f"{{{A}}}ext").attrib.update({"cx": str(cx), "cy": str(cy)})
    
    prstGeom = ET.SubElement(spPr, f"{{{A}}}prstGeom")
    prstGeom.set("prst", "rect")
    ET.SubElement(prstGeom, f"{{{A}}}avLst")
    ln = ET.SubElement(spPr, f"{{{A}}}ln")
    ln.set("w", "0")
    ET.SubElement(ln, f"{{{A}}}noFill")
    
    body.insert(0, p)
    tree.write(doc_xml, xml_declaration=True, encoding="UTF-8", standalone=True)
    print("DEBUG: [insert_cover_image] Finished.")





def color_list_markers(temp_dir: str, accent_color: str):
    """
    Surgically injects ACCENT_COLOR into word/numbering.xml to color bullets and numbers.
    """
    num_xml_path = os.path.join(temp_dir, "word", "numbering.xml")
    if not os.path.exists(num_xml_path):
        return
    
    try:
        tree = ET.parse(num_xml_path)
        root = tree.getroot()
        changed = False
        
        # We target all <w:lvl> elements (levels of a list)
        for lvl in root.findall(f".//{{{W}}}lvl"):
            rPr = lvl.find(f"{{{W}}}rPr")
            if rPr is None:
                # Insert rPr before the first non-pPr child if possible, or just append
                rPr = ET.Element(f"{{{W}}}rPr")
                lvl.append(rPr)
            
            # Set or update color
            color = rPr.find(f"{{{W}}}color")
            if color is None:
                color = ET.SubElement(rPr, f"{{{W}}}color")
            color.set(f"{{{W}}}val", accent_color)
            
            # Set or update Font (Cambria)
            rFonts = rPr.find(f"{{{W}}}rFonts")
            if rFonts is None:
                rFonts = ET.SubElement(rPr, f"{{{W}}}rFonts")
            rFonts.set(f"{{{W}}}ascii", FONT_NAME)
            rFonts.set(f"{{{W}}}hAnsi", FONT_NAME)
            rFonts.set(f"{{{W}}}cs", FONT_NAME)
            rFonts.set(f"{{{W}}}eastAsia", FONT_NAME)
            
            changed = True
            
        if changed:
            tree.write(num_xml_path, xml_declaration=True, encoding="UTF-8", standalone=True)
    except Exception as e:
        print(f"Error coloring list markers: {e}")

# ─── Low-level XML Pipeline ──────────────────────────────────────────────────

def low_level_xml_processing(docx_path: str, output_path: str, template_path: str = None, cover_image_path: str = None, theme=None):
    temp = f"xml_tmp_{uuid.uuid4().hex[:8]}"
    os.makedirs(temp, exist_ok=True)
    with zipfile.ZipFile(docx_path, "r") as z: z.extractall(temp)

    
    if template_path: inject_headers_footers(template_path, temp)
    
    apply_image_borders_and_centering(temp, theme)
    tune_hf_xml(temp, theme)
    
    # 25% Sharpness via Pixel Manipulation (Pillow) - EXCLUDING COVER
    apply_pixel_sharpness(temp)
    
    if cover_image_path: insert_cover_image(temp, cover_image_path)
    
    # ─── Color and Format List Markers ───
    accent_color = (theme.get('accent') or theme.get('headerFooter') or ACCENT_COLOR).replace("#", "")
    color_list_markers(temp, accent_color)
    
    # Final check on Content Types BEFORE repack
    ensure_content_types(temp)



    
    # Repack
    with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for r_dir, dirs, files in os.walk(temp):
            for file in files:
                fp = os.path.join(r_dir, file)
                rel_p = os.path.relpath(fp, temp).replace("\\", "/")
                zout.write(fp, rel_p)
    
    safe_rmtree(temp)



# ─── HTML to DOCX Conversion ──────────────────────────────────────────────────

@app.post("/convert-message")
async def convert_message(
    html: str = Form(...),
    filename: str = Form("documento.docx")
):
    try:
        working_dir = f"temp_conv_{uuid.uuid4().hex}"
        os.makedirs(working_dir, exist_ok=True)
        out_f = os.path.join(working_dir, "output.docx")
        
        # Create a new document
        doc = Document()
        
        # Set Page Size to A4 (21.0 x 29.7 cm)
        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        
        # Set Margins (Top 3, Bottom 2, Left 3, Right 2)
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        
        # Convert HTML to Docx
        new_parser = HtmlToDocx()
        new_parser.add_html_to_document(html, doc)
        
        # Apply some basic formatting if needed
        for para in doc.paragraphs:
            if not para.style:
                para.style = doc.styles['Normal']
            for run in para.runs:
                run.font.name = FONT_NAME
                run.font.size = Pt(11)

        doc.save(out_f)
        
        # Response path
        response_path = f"final_msg_{uuid.uuid4().hex}.docx"
        shutil.copy(out_f, response_path)
        shutil.rmtree(working_dir)
        
        return FileResponse(
            response_path, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
            filename=filename
        )
    except Exception as e:
        print(f"Error converting HTML to Docx: {e}")
        return {"error": str(e)}

# ─── Main Endpoint ───────────────────────────────────────────────────────────

@app.post("/format")
async def format_docx(
    target_docs: List[UploadFile] = File(...),
    template_docx: UploadFile = File(None),
    theme: str = Form(None),
    indent: str = Form("true")
):
    indent_flag = indent.lower() == "true"
    theme_data = json.loads(theme) if theme else None
    
    # Setup template
    # Dynamic path detection for both Windows Local and Linux VPS
    if os.name == 'nt':
        static_tpl = r"C:\poisson-erp\public\covers\Modelo.docx"
    else:
        static_tpl = "/var/www/poisson-erp/public/covers/Modelo.docx"
    
    working_dir = f"temp_{uuid.uuid4().hex}"
    os.makedirs(working_dir, exist_ok=True)
    tpl_p = os.path.join(working_dir, "template.docx")
    
    if template_docx:
        with open(tpl_p, "wb") as f: f.write(await template_docx.read())
    else:
        if os.path.exists(static_tpl):
            shutil.copy(static_tpl, tpl_p)
        else:
            # Fallback relative to sibling directory
            local_tpl = os.path.join(os.path.dirname(__file__), "..", "poisson-erp", "public", "covers", "Modelo.docx")
            if os.path.exists(local_tpl):
                shutil.copy(local_tpl, tpl_p)
            else:
                return {"error": f"Arquivo Modelo.docx não encontrado no servidor. (Tentado: {static_tpl})"}

    # Resolve Cover Image
    cover_img = None
    if theme_data and 'image' in theme_data:
        img_rel = theme_data['image'].lstrip('/')
        if os.name == 'nt':
            base_public = r"C:\poisson-erp\public"
        else:
            base_public = "/var/www/poisson-erp/public"
            
        p1 = os.path.join(base_public, img_rel.replace("/", os.sep))
        p2 = os.path.join(os.path.dirname(__file__), "..", "poisson-erp", "public", img_rel.replace("/", os.sep))
        
        if os.path.exists(p1): cover_img = p1
        elif os.path.exists(p2): cover_img = p2
    
    if not cover_img:
        cover_img = extract_first_page_image(tpl_p)
    
    print(f"DEBUG: [format_docx] Final resolved cover_img: {cover_img}")

    processed_files = []

    for idx, target_docx in enumerate(target_docs):
        batch_id = f"file_{idx}_{uuid.uuid4().hex[:6]}"
        tgt_p = os.path.join(working_dir, f"{batch_id}_in.docx")
        inter_p = os.path.join(working_dir, f"{batch_id}_inter.docx")
        out_f = os.path.join(working_dir, f"{batch_id}_out.docx")
        
        # Save input
        content = await target_docx.read()
        with open(tgt_p, "wb") as f: f.write(content)
        
        # PDF to DOCX Conversion if needed
        if target_docx.filename.lower().endswith(".pdf"):
            docx_from_pdf = os.path.join(working_dir, f"{batch_id}_conv.docx")
            try:
                cv = Converter(tgt_p)
                cv.convert(docx_from_pdf, start=0, end=None)
                cv.close()
                tgt_p = docx_from_pdf # Switch to the converted docx
            except Exception as e:
                print(f"Erro ao converter PDF {target_docx.filename}: {e}")
                # Fallback or error? For now, we continue and hope Document(tgt_p) catches it or we skip
        
        doc = Document(tgt_p)

        
        # Formatting Pipeline
        remove_section_breaks(doc)
        insert_chapter_title(doc, target_docx.filename, theme_data)
        apply_global_styles(doc, theme_data, indent_flag=indent_flag)
        detect_and_format_titles(doc.paragraphs, theme_data)
        remove_empty_lines_after_titles(doc.paragraphs)
        format_tables(doc, theme_data)
        
        # Final section break requested only at the end
        
        kws = ["palavras-chave", "palavra-chave", "palavras chave"]
        for para in doc.paragraphs:
            if any(k in para.text.lower() for k in kws): para.add_run().add_break(WD_BREAK.PAGE)
        
        doc.add_section() # Single break at the end
        doc.save(inter_p)
        
        # Low-level processing
        low_level_xml_processing(inter_p, out_f, template_path=tpl_p, cover_image_path=cover_img, theme=theme_data)
        
        # Determine internal-zip filename
        orig_name = target_docx.filename or f"documento_{idx}"
        base_name, _ = os.path.splitext(orig_name)
        final_name_in_zip = f"{base_name}_formatado.docx"
        
        processed_files.append((out_f, final_name_in_zip))

    # Response Logic
    if len(processed_files) == 1:
        out_path, final_filename = processed_files[0]
        # Copy to a stable location for FileResponse
        response_path = f"final_{uuid.uuid4().hex}.docx"
        shutil.copy(out_path, response_path)
        shutil.rmtree(working_dir)
        return FileResponse(response_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=final_filename)
    else:
        zip_path = f"batch_{uuid.uuid4().hex}.zip"
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for out_path, final_name in processed_files:
                zout.write(out_path, final_name)
        
        shutil.rmtree(working_dir)
        return FileResponse(zip_path, media_type="application/zip", filename="documentos_formatados.zip")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8030)

